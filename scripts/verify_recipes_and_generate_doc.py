# -*- coding: utf-8 -*-
"""核验食谱营养、补全步骤，并生成完整食谱数据手册。"""
from __future__ import annotations

import re
import shutil
import sqlite3
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DB = ROOT / "data" / "diet.db"
OUT_DIR = ROOT / "食谱数据"
OUT_DOCX = OUT_DIR / "智能营养膳食推荐系统_食谱数据手册.docx"
BACKUP = OUT_DIR / "diet_before_recipe_verification.db"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "E8EEF5"
GRAY = "5F6B7A"
GRID = "C9D3DF"


def normalize_steps(text: str, name: str, role: str, ingredients: list[dict]) -> str:
    """保留有效原步骤；不完整时按菜品角色生成至少 4 步可执行说明。"""
    text = (text or "").replace("\r", "\n").strip()
    lines = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line or line.startswith("用料："):
            continue
        line = re.sub(r"^\s*\d+[\.、)）]\s*", "", line).strip()
        if line and line not in lines:
            lines.append(line)

    vague = any(x in text for x in ("详见原文", "暂无步骤", "适量调味即可"))
    if len(lines) >= 3 and not vague:
        if len(lines) == 3:
            lines.append("装盘后趁热食用；冷食或饮品按菜品特点冷藏后饮用。")
        return "\n".join(f"{i}. {line}" for i, line in enumerate(lines, 1))

    names = [x["name"] for x in ingredients]
    main = "、".join(names[:5]) or "所需食材"
    if role == "drink":
        generated = [
            f"准备{main}；水果洗净去皮去核，乳品或豆浆保持冷藏。",
            "将固体食材切成小块，与适量饮用水或乳品一起放入料理机。",
            "搅打至细腻均匀；需要温热的豆浆或牛奶用小火加热，避免沸腾溢出。",
            "按个人口味调整浓稠度，不额外添加糖或仅少量添加，倒杯后尽快饮用。",
        ]
    elif role == "staple":
        generated = [
            f"称取{main}，淘洗 1-2 次后沥水。",
            "按米与水约 1:1.2 的比例加水，浸泡 15-20 分钟。",
            "放入电饭煲按标准煮饭程序加热，结束后焖 8-10 分钟。",
            "用饭勺翻松，按推荐份量盛出，避免额外拌油。",
        ]
    elif role == "soup":
        generated = [
            f"将{main}洗净；肉类切块并焯水，蔬菜切成大小均匀的块。",
            "锅中加入清水，先放耐煮的肉类或根茎食材，大火煮沸后撇去浮沫。",
            "转小火煮至主料熟透，再放入易熟蔬菜，继续煮 5-10 分钟。",
            "出锅前少量加盐调味，避免重复加入高钠酱料，盛碗后趁热食用。",
        ]
    elif role == "vegetable":
        generated = [
            f"将{main}清洗干净，切成大小均匀的片、段或块并沥干。",
            "炒锅预热后加入配方中的食用油，先下葱姜蒜等香辛料小火炒香。",
            "转大火放入蔬菜快速翻炒；质地较硬的食材可先焯水 1-2 分钟。",
            "炒至断生后少量加盐或酱油调味，立即出锅以减少营养流失。",
        ]
    else:
        generated = [
            f"将{main}按用量备齐；肉类切块或切片，蔬菜洗净切配，生熟食材分开处理。",
            "肉类用少量酱油、淀粉或姜片腌制 10 分钟；需要焯水的食材先焯水沥干。",
            "锅中加入配方中的食用油，先将肉类煎炒至变色，再加入蔬菜或辅料翻炒。",
            "加入少量清水或调味汁焖至食材完全熟透，收汁后少量加盐调味并装盘。",
        ]
    return "\n".join(f"{i}. {line}" for i, line in enumerate(generated, 1))


def find_food(conn: sqlite3.Connection, keywords: list[str]) -> tuple | None:
    for keyword in keywords:
        row = conn.execute(
            "SELECT id,name,calories,protein,fat,carbs FROM foods "
            "WHERE name LIKE ? AND calories IS NOT NULL ORDER BY length(name) LIMIT 1",
            (f"%{keyword}%",),
        ).fetchone()
        if row:
            return row
    return None


def ensure_recipe_food_links(conn: sqlite3.Connection) -> None:
    """给标准白米饭与种子饮品建立明确的食材关系。"""
    rice_food = find_food(conn, ["米饭(蒸)(均值)", "米饭"])
    if not rice_food:
        raise RuntimeError("foods 表中未找到米饭营养数据")

    rice = conn.execute("SELECT id FROM recipes WHERE name='白米饭' LIMIT 1").fetchone()
    if rice:
        rice_id = rice[0]
        conn.execute(
            "UPDATE recipes SET category='午餐', dish_role='staple', cook_minutes=30, "
            "source='verified_seed' WHERE id=?", (rice_id,)
        )
    else:
        cur = conn.execute(
            "INSERT INTO recipes(name,category,steps,cook_minutes,accent,dish_role,source) "
            "VALUES('白米饭','午餐','',30,'green','staple','verified_seed')"
        )
        rice_id = cur.lastrowid
    conn.execute("DELETE FROM recipe_foods WHERE recipe_id=?", (rice_id,))
    conn.execute(
        "INSERT INTO recipe_foods(recipe_id,food_id,quantity) VALUES(?,?,?)",
        (rice_id, rice_food[0], 150.0),
    )

    drink_aliases = {
        "苹果汁": ["苹果汁", "苹果"], "梨子汁": ["梨汁", "梨"],
        "橙汁": ["橙汁", "橙"], "牛奶": ["牛乳", "牛奶"],
        "豆浆": ["豆浆", "豆奶"], "酸奶": ["酸奶", "酸牛奶"],
    }
    for name, aliases in drink_aliases.items():
        row = conn.execute("SELECT id FROM recipes WHERE name=?", (name,)).fetchone()
        if not row:
            continue
        rid = row[0]
        if conn.execute("SELECT 1 FROM recipe_foods WHERE recipe_id=?", (rid,)).fetchone():
            continue
        food = find_food(conn, aliases)
        if food:
            conn.execute(
                "INSERT OR REPLACE INTO recipe_foods(recipe_id,food_id,quantity) VALUES(?,?,?)",
                (rid, food[0], 200.0),
            )


def verify_all_recipes(conn: sqlite3.Connection) -> tuple[list[dict], dict]:
    recipes = []
    stats = {"verified": 0, "unlinked": 0, "steps_completed": 0}
    rows = conn.execute(
        "SELECT id,name,category,COALESCE(dish_role,'mixed'),COALESCE(steps,''),"
        "COALESCE(cook_minutes,20),COALESCE(source,'本地'),COALESCE(source_ref,''),"
        "COALESCE(total_weight,0),per100_calories,per100_protein,per100_fat,per100_carbs,"
        "COALESCE(nutrition_verified_at,'') FROM recipes ORDER BY id"
    ).fetchall()
    for (rid, name, category, role, steps, minutes, source, source_ref, total_weight,
         p100_cal, p100_protein, p100_fat, p100_carbs, verified_status) in rows:
        ing_rows = conn.execute(
            "SELECT f.id,COALESCE(NULLIF(rf.display_name,''),f.name),rf.quantity,"
            "f.calories,f.protein,f.fat,f.carbs,COALESCE(f.unit,'100g'),COALESCE(f.source,''),"
            "COALESCE(rf.source_text,'') "
            "FROM recipe_foods rf JOIN foods f ON f.id=rf.food_id "
            "WHERE rf.recipe_id=? ORDER BY rf.rowid", (rid,)
        ).fetchall()
        ingredients = []
        cal = protein = fat = carbs = 0.0
        nutrition_complete = not verified_status.startswith("待精确检索")
        for fid, food_name, qty, fcal, fp, ff, fc, unit, food_source, source_text in ing_rows:
            factor = float(qty) / 100.0
            ignored_micro = "忽略营养贡献" in food_source or "【营养忽略】" in source_text
            contribution = {
                "calories": 0.0 if ignored_micro else (None if fcal is None else fcal * factor),
                "protein": 0.0 if ignored_micro else (None if fp is None else fp * factor),
                "fat": 0.0 if ignored_micro else (None if ff is None else ff * factor),
                "carbs": 0.0 if ignored_micro else (None if fc is None else fc * factor),
            }
            cal += contribution["calories"] or 0.0
            protein += contribution["protein"] or 0.0
            fat += contribution["fat"] or 0.0
            carbs += contribution["carbs"] or 0.0
            ingredients.append({
                "id": fid, "name": food_name, "quantity": float(qty), "unit": "g",
                **contribution, "source": food_source, "ignored_micro": ignored_micro,
            })
        if ingredients:
            stats["verified"] += 1
        else:
            stats["unlinked"] += 1

        if not ingredients:
            continue
        nutrition_complete = nutrition_complete and None not in (p100_cal, p100_protein, p100_fat, p100_carbs)
        if nutrition_complete and abs(cal - p100_cal * total_weight / 100.0) > 0.2:
            raise RuntimeError(f"食谱 {rid}/{name} 的每100g与整份热量不一致")
        if not nutrition_complete:
            stats["unlinked"] += 1
        recipes.append({
            "id": rid, "name": name, "category": category, "role": role,
            "steps": steps, "minutes": minutes, "source": source, "source_ref": source_ref,
            "ingredients": ingredients,
            "calories": round(cal, 1) if nutrition_complete else None,
            "protein": round(protein, 1) if nutrition_complete else None,
            "fat": round(fat, 1) if nutrition_complete else None,
            "carbs": round(carbs, 1) if nutrition_complete else None,
            "total_weight": round(total_weight, 1),
            "per100_calories": round(p100_cal, 1) if nutrition_complete else None,
            "per100_protein": round(p100_protein, 1) if nutrition_complete else None,
            "per100_fat": round(p100_fat, 1) if nutrition_complete else None,
            "per100_carbs": round(p100_carbs, 1) if nutrition_complete else None,
            "nutrition_complete": nutrition_complete, "verified_status": verified_status,
        })
    return recipes, stats


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Microsoft YaHei", size=10.5, color="000000", bold=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("智能营养膳食推荐系统 | 食谱数据手册"), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("第 "), size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    set_font(footer.add_run(" 页"), size=9, color=GRAY)


def add_recipe(doc: Document, recipe: dict, index: int) -> None:
    h = doc.add_heading(f"{index}. {recipe['name']}", level=2)
    h.paragraph_format.page_break_before = index > 1

    meta = doc.add_table(rows=4, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(meta, [1500, 3180, 1500, 3180])
    if recipe["nutrition_complete"]:
        per100_text = f"{recipe['per100_calories']:.1f} kcal；蛋白质 {recipe['per100_protein']:.1f} g；脂肪 {recipe['per100_fat']:.1f} g；碳水 {recipe['per100_carbs']:.1f} g"
        total_text = f"{recipe['calories']:.1f} kcal；蛋白质 {recipe['protein']:.1f} g"
        macro_text = f"脂肪 {recipe['fat']:.1f} g；碳水 {recipe['carbs']:.1f} g"
    else:
        per100_text = total_text = macro_text = "待精确核验（存在未取得同物种营养数据的食材，不输出近似总值）"
    values = [
        ("餐次/类型", f"{recipe['category']} / {recipe['role']}"),
        ("烹饪时间", f"约 {recipe['minutes']} 分钟"),
        ("每100g", per100_text),
        ("整份重量", f"约 {recipe['total_weight']:.1f} g"),
        ("整份营养", total_text),
        ("整份宏量", macro_text),
        ("数据来源", recipe['source']),
        ("来源编号", recipe['source_ref'] or "标准化基础食谱"),
    ]
    for pos, (label, value) in enumerate(values):
        r, c = divmod(pos * 2, 4)
        meta.cell(r, c).text = label
        meta.cell(r, c + 1).text = value
        set_cell_shading(meta.cell(r, c), LIGHT_FILL)
        for run in meta.cell(r, c).paragraphs[0].runs:
            set_font(run, size=9, color=DARK_BLUE, bold=True)
        for run in meta.cell(r, c + 1).paragraphs[0].runs:
            set_font(run, size=9)

    doc.add_heading("用料与营养贡献", level=3)
    ing = doc.add_table(rows=1, cols=6)
    ing.style = "Table Grid"
    ing.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [2160, 1050, 1500, 1500, 1500, 1650]
    set_table_geometry(ing, widths)
    headers = ["食材", "用量", "热量(kcal)", "蛋白质(g)", "脂肪(g)", "碳水(g)"]
    for idx, text in enumerate(headers):
        ing.cell(0, idx).text = text
        set_cell_shading(ing.cell(0, idx), LIGHT_FILL)
        for run in ing.cell(0, idx).paragraphs[0].runs:
            set_font(run, size=8.5, color=DARK_BLUE, bold=True)
    if recipe["ingredients"]:
        for item in recipe["ingredients"]:
            cells = ing.add_row().cells
            fmt = lambda value: "待核验" if value is None else f"{value:.1f}"
            if item.get("ignored_micro"):
                nutrient_cells = ["忽略", "忽略", "忽略", "忽略"]
            else:
                nutrient_cells = [fmt(item["calories"]), fmt(item["protein"]), fmt(item["fat"]), fmt(item["carbs"])]
            vals = [item["name"], f"{item['quantity']:.1f} g", *nutrient_cells]
            for idx, value in enumerate(vals):
                cells[idx].text = value
                for run in cells[idx].paragraphs[0].runs:
                    set_font(run, size=8.5)
        set_table_geometry(ing, widths)
    else:
        cells = ing.add_row().cells
        cells[0].merge(cells[-1]).text = "暂无可关联食材，营养值按 0 计；需后续人工补录。"
        set_table_geometry(ing, widths)

    doc.add_heading("制作步骤", level=3)
    for line in recipe["steps"].split("\n"):
        p = doc.add_paragraph(line.strip())
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_together = True
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.keep_with_next = False
    set_font(note.add_run("核验口径："), size=8.5, color=GRAY, bold=True)
    set_font(note.add_run("整份营养由食材用量 × foods 每100g营养值累加；每100g营养再按整份总重量换算。≤12g的盐、鸡精及香辛料按微量调料忽略营养贡献，均四舍五入到0.1。"), size=8.5, color=GRAY)


def build_document(recipes: list[dict], stats: dict) -> None:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("智能营养膳食推荐系统"), size=14, color=GRAY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_font(p.add_run("食谱数据手册"), size=28, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("用料 · 制作步骤 · 营养核验"), size=13, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    set_font(p.add_run(f"共 {len(recipes)} 道食谱 | 生成日期：{datetime.now():%Y-%m-%d}"), size=10, color=GRAY)
    doc.add_page_break()

    doc.add_heading("数据说明", level=1)
    for text in (
        f"本手册收录重建数据库中的 {len(recipes)} 道食谱。所有来源原料均进入用料表；其中 {len(recipes)-stats['unlinked']} 道已完成营养复算，{stats['unlinked']} 道保留菜谱并标记待检索。",
        "营养匹配顺序为：优先使用中国食物成分表中的精确食材；无精确项时使用明确的同类别食材估算（例如鸡腿可回退到鸡肉均值），并保留原始食材名称以便追溯。",
        "结构化食谱来自“菜谱数据库.mdb”；《中国名菜谱》17册作为菜名、地方风味和传统做法的参考资料。扫描版PDF不具备文本层，因此仅对可人工辨认内容作辅助校对。",
        "鸡蛋统一按全蛋平均值估算，每个约50g；酸奶和牛奶采用无品牌均值。食材展示名按日常名称规范化，但底层仍保留原始来源字段。",
        "午餐和晚餐推荐固定包含一份标准白米饭（150g），且主食不重复添加；荤菜、素菜和汤按照推荐引擎另行组合。",
        "食谱仅用于课程项目与一般饮食参考，不构成医疗建议。",
    ):
        doc.add_paragraph(text)
    doc.add_heading("营养计算公式", level=1)
    doc.add_paragraph("单项营养贡献 = 食材用量(g) ÷ 100 × 食材每100g营养值；食谱总营养 = 各食材贡献之和。")
    doc.add_paragraph("≤12g的盐、鸡精、大料、八角、花椒等微量调味料或香辛料不计营养贡献；烹饪损耗和品牌差异仍会造成实际值波动。")

    for index, recipe in enumerate(recipes, 1):
        add_recipe(doc, recipe, index)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB)
    conn.execute("PRAGMA foreign_keys=ON")
    recipes, stats = verify_all_recipes(conn)
    conn.close()
    build_document(recipes, stats)
    print({"document": str(OUT_DOCX), "recipes": len(recipes), **stats})


if __name__ == "__main__":
    main()
