# -*- coding: utf-8 -*-
"""Generate Project Start Report and Research Report as Word documents."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

BASE = r"C:\Users\ROG\Documents\System"
CONTENT = os.path.join(BASE, "report_content")
STUDENT_ID = "20252323"
STUDENT_NAME = "宋启铖"
AUTHOR = STUDENT_NAME


def read_txt(name):
    with open(os.path.join(CONTENT, name), encoding="utf-8") as f:
        return f.read().strip()


def set_doc_font(doc, font_name="宋体", size=12):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    return p


def add_multiline(doc, text):
    for line in text.split("\n"):
        if line.strip():
            add_para(doc, line.strip())


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def build_start_report():
    doc = Document()
    set_doc_font(doc)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("智能营养膳食推荐系统项目开题报告")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()
    cover = [
        ("项目名称", "智能营养膳食推荐系统"),
        ("密级", "仅供收件方查阅"),
        ("项目编号", "NDS-2026-001"),
        ("版本", "V1.0"),
        ("文档编号", "Project ID_INIT_002"),
        ("拟制", AUTHOR),
        ("拟制日期", "2026-08-27"),
        ("评审日期", "2026-08-28"),
        ("批准日期", "2026-08-29"),
    ]
    for k, v in cover:
        add_para(doc, f"{k}：{v}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("武汉学链科技有限公司").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("版权所有  不得复制")

    doc.add_page_break()

    add_heading(doc, "Revision Record（修订记录）", 2)
    add_table(
        doc,
        ["Date 日期", "Revision Version", "CR ID", "Sec No.", "Change Description", "Author"],
        [
            [
                "2026-08-27",
                "V1.0",
                "INIT-001",
                "全部章节",
                "基于MVP创建初始版本，完成需求分析、原型设计、任务拆解与AI协作规划",
                AUTHOR,
            ]
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Catalog（目录）", 2)
    for item in [
        "1  项目提出",
        "2  开发团队组成和计划时间",
        "3  项目预计支出",
        "4  风险评估和规避（含AI协作专项）",
    ]:
        add_para(doc, item)

    doc.add_page_break()

    # Section 1
    add_heading(doc, "1. 项目提出", 1)
    add_heading(doc, "1.1 项目名称", 2)
    add_para(doc, "智能营养膳食推荐系统（Smart Diet Recommendation System）")

    add_heading(doc, "1.2 项目简介", 2)
    add_multiline(doc, read_txt("section1_intro.txt").split("【需求分析与用户画像】")[0].strip())

    add_heading(doc, "1.3 项目目标", 2)
    add_multiline(doc, read_txt("section1_goals.txt"))

    add_heading(doc, "1.4 系统边界", 2)
    add_multiline(doc, read_txt("section1_boundary.txt"))

    add_heading(doc, "1.5 需求分析与用户画像", 2)
    intro_full = read_txt("section1_intro.txt")
    if "【需求分析与用户画像】" in intro_full:
        add_multiline(doc, intro_full.split("【需求分析与用户画像】")[1])

    add_heading(doc, "1.6 功能矩阵与优先级（MoSCoW）", 2)
    add_table(
        doc,
        ["需求ID", "模块", "功能点", "优先级", "计划完成时间"],
        [
            ["F1", "用户", "注册/登录（本地账号）", "P0", "Day 1-2"],
            ["F2", "用户", "身高/体重录入与热量目标自动计算", "P0", "Day 2"],
            ["F3", "数据", "预置100+食材与30+食谱初始化", "P0", "Day 1"],
            ["F4", "推荐", "基于目标的智能三餐生成（规则引擎）", "P0", "Day 4-5"],
            ["F5", "展示", "三餐卡片展示（含热量/蛋白质/碳水/脂肪）", "P0", "Day 4"],
            ["F6", "展示", "食谱详情弹窗（食材清单+制作步骤）", "P0", "Day 5"],
            ["F7", "食材", "按名称搜索食材库", "P1", "Day 6"],
            ["F8", "用户", "偏好设置（如不吃辣）", "P2", "视进度而定"],
            ["F9", "推荐", "换一餐替换功能", "P2", "V2.0"],
        ],
    )

    add_heading(doc, "1.7 原型设计（UI低保真蓝图）", 2)
    add_para(doc, "主窗口 900×700：顶部导航（首页/食材库/智能推荐），中部展示目标、热量预算与BMI；三餐卡片按早30%、午40%、晚30%分配热量；支持重新生成方案；食谱详情弹窗展示食材克数与制作步骤。")

    add_heading(doc, "1.8 工作量估计", 2)
    workload = [
        ("环境搭建与数据库设计", "Qt配置、SQLite建表、初始化脚本", "0.8", "含预置100条食材+30条食谱"),
        ("用户管理模块", "注册登录、目标设置与BMI计算", "1.2", "连接数据库存储用户记录"),
        ("数据展示模块", "食材库列表、搜索框、食谱卡片UI", "1.0", "QListWidget与QLineEdit"),
        ("智能推荐引擎", "营养规则算法、热量分配器", "1.5", "核心算法，人工精调系数"),
        ("推荐结果展示", "三餐卡片动态刷新、营养标签", "1.0", "布局管理、信号槽绑定"),
        ("食谱详情弹窗", "模态对话框、食材与步骤渲染", "0.8", "QTextEdit或QLabel排版"),
        ("系统集成与界面打磨", "导航切换、窗口自适应", "0.5", "提升用户体验"),
        ("测试与Bug修复", "端到端测试、边界Case", "0.7", "极端身高体重、空数据等"),
        ("文档与答辩准备", "开题报告、演示脚本", "0.5", "整理AI协作记录"),
        ("总工作量（人天）", "合计", "8.0", "7自然日交付，预留1天缓冲"),
    ]
    add_table(doc, ["模块", "子模块", "工作量估计（人天）", "说明"], workload)
    add_para(doc, "备注：“人天”即1个人工作8小时的量。本计划按高强度冲刺排期，预留1天缓冲。")

    doc.add_page_break()

    # Section 2
    add_heading(doc, "2. 开发团队组成和计划时间", 1)
    add_para(doc, "项目周期：2026年08月27日 - 2026年09月02日（计7天）")
    add_para(doc, "项目总监：1人  姓名：王文鑫")
    add_para(doc, f"项目经理/核心开发：1人  姓名：{AUTHOR}")
    add_para(doc, "项目成员：1人（辅助测试与数据校验）")
    add_para(doc, "人员来源：重庆大学2025级大数据与软件学院")
    add_heading(doc, "7日冲刺计划", 2)
    add_multiline(doc, read_txt("section2_schedule.txt"))

    add_table(
        doc,
        ["天数", "冲刺目标", "关键交付物", "验收标准"],
        [
            ["Day 1", "环境与数据基石", "可运行工程+SQLite数据", "foods表≥100条"],
            ["Day 2", "用户入口闭环", "注册登录、目标设置", "可保存用户与热量目标"],
            ["Day 3", "食谱与食材可视化", "食材列表、食谱卡片", "展示30道菜名卡片"],
            ["Day 4", "核心算法攻坚", "RecommendEngine", "热量±5%输出3道菜"],
            ["Day 5", "全流程打通", "三餐推荐+详情弹窗", "完整走通主流程"],
            ["Day 6", "集成与美化", "搜索、界面微调", "搜索鸡过滤鸡肉食材"],
            ["Day 7", "验收与封板", "exe、演示视频、PPT", "试运行无崩溃"],
        ],
    )

    doc.add_page_break()

    # Section 3
    add_heading(doc, "3. 项目预计支出", 1)
    items = [
        ("设备、场地占用费", "无"),
        ("本地人员工资（管理费）", "无（教育项目，不计人力成本）"),
        ("外协人员工资", "无"),
        ("加班费", "无"),
        ("交通费", "无"),
        ("住宿费", "无"),
        ("其它费用", "无"),
        ("总计", "0 元"),
    ]
    for k, v in items:
        add_para(doc, f"{k}：{v}")
    add_para(doc, read_txt("section3_note.txt"))

    doc.add_page_break()

    # Section 4
    add_heading(doc, "4. 风险评估和规避（含AI协作专项）", 1)
    add_heading(doc, "4.1 技术风险", 2)
    add_multiline(doc, read_txt("section4_tech.txt"))
    add_heading(doc, "4.2 管理风险", 2)
    add_multiline(doc, read_txt("section4_mgmt.txt"))
    add_heading(doc, "4.3 其它风险", 2)
    add_multiline(doc, read_txt("section4_other.txt"))

    add_heading(doc, "4.4 AI协作过程量化记录", 2)
    add_table(
        doc,
        ["迭代轮次", "我的输入", "AI产出", "批判性修正"],
        [
            ["Round 1", "SQLite初始化表代码", "标准建表语句", "calories改REAL，增加unit字段"],
            ["Round 2", "增肌减脂热量公式", "Harris-Benedict公式", "加入性别参数，微调活动系数"],
            ["Round 3", "推荐算法伪代码", "机器学习协同过滤", "否决，改为贪心+约束校验O(n)"],
            ["Round 4", "LNK2019报错", "添加QT+=sql", "确认.pro模块，统一MinGW编译"],
        ],
    )

    add_para(doc, "总结：将AI定位于超级实习生，用于重复性编码与信息检索；业务决策、架构设计与错误边界处理由开发者主导。")

    out = os.path.join(BASE, f"{STUDENT_ID}_{STUDENT_NAME}_Project Start Report_V1.0.docx")
    doc.save(out)
    return out


def build_research_report():
    doc = Document()
    set_doc_font(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("智能营养膳食推荐系统 — 调研报告")
    r.bold = True
    r.font.size = Pt(22)

    add_para(doc, "项目名称：智能营养膳食推荐系统")
    add_para(doc, "文档版本：V1.0")
    add_para(doc, "编制日期：2026-08-27")
    add_para(doc, f"编制人：{AUTHOR}")
    add_para(doc, "所属单位：重庆大学2025级大数据与软件学院")

    doc.add_page_break()

    add_heading(doc, "1. 调研背景与目的", 1)
    add_multiline(
        doc,
        """
随着居民健康意识提升，个性化膳食管理需求增长，但现有工具普遍存在中式菜谱覆盖不足、推荐算法不透明、依赖网络等问题。
本调研旨在：分析目标用户痛点与使用场景；对比主流竞品能力差距；验证基于规则引擎的本地化桌面方案可行性；为7日MVP开发提供需求与设计依据。
        """.strip(),
    )

    add_heading(doc, "2. 市场调研", 1)
    add_heading(doc, "2.1 行业现状", 2)
    add_multiline(
        doc,
        """
健康管理类应用市场规模持续扩大，饮食记录与营养分析是高频功能。移动端App占主导，但桌面端在数据隐私、离线使用、大屏信息展示方面仍有空间。
专业用户（健身增肌、科学减脂）更关注宏量营养素达标，而非单纯热量记账。
        """.strip(),
    )

    add_heading(doc, "2.2 目标用户调研", 2)
    add_table(
        doc,
        ["维度", "描述"],
        [
            ["用户代号", "张明，25岁，重庆互联网公司程序员"],
            ["生活特征", "独居、加班、工作日外卖为主、晚餐偶尔自炊"],
            ["健康目标", "3个月内增肌3kg，日均蛋白质目标≥100g"],
            ["痛点", "外卖油腻营养不均；自炊不知吃什么；现有App中式菜谱少"],
            ["使用场景", "晚间8点设置目标，查看次日三餐与采购清单"],
            ["偏好", "中式家常菜、口感优先、操作≤5分钟"],
        ],
    )

    add_heading(doc, "3. 竞品分析", 1)
    add_table(
        doc,
        ["竞品", "优势", "劣势/差距", "本项目机会"],
        [
            ["MyFitnessPal", "全球最大食材库，记录精细", "难自动生成三餐，中式菜谱少", "本地化菜谱+一键生成"],
            ["薄荷健康", "中文社区成熟，食物库丰富", "广告绑定、算法不透明、需联网", "规则透明+完全离线"],
            ["ChatGPT/豆包", "菜谱生成灵活", "营养总量不可控，建议荒谬", "强制营养校验±5%"],
            ["Keep饮食模块", "与运动场景结合", "膳食推荐非核心，深度不足", "专注膳食规划垂直场景"],
        ],
    )

    add_heading(doc, "4. 技术调研", 1)
    add_heading(doc, "4.1 开发技术选型", 2)
    add_table(
        doc,
        ["技术", "选型", "理由"],
        [
            ["UI框架", "Qt 6 + C++", "课程要求、跨平台、原生桌面体验"],
            ["数据库", "SQLite", "轻量、嵌入式、无需独立服务"],
            ["推荐算法", "规则引擎+贪心约束校验", "可解释、O(n)复杂度、适合MVP"],
            ["营养计算", "食材关联表累加", "避免AI/手工填入总营养造成幻觉"],
        ],
    )

    add_heading(doc, "4.2 营养学规则参考", 2)
    add_multiline(
        doc,
        """
参考《中国居民膳食指南（2022）》及运动营养常识：
• 三餐热量分配建议：早餐30%、午餐40%、晚餐30%。
• 增肌期：蛋白质摄入建议≥1.8g/kg体重。
• 减脂期：碳水供能比建议≤55%。
• 基础代谢：采用 Harris-Benedict 公式，结合性别与活动系数修正。
        """.strip(),
    )

    add_heading(doc, "5. 需求调研结论", 1)
    add_multiline(
        doc,
        """
经调研，MVP应聚焦以下核心闭环：
1. 用户录入身体数据与健康目标；
2. 系统基于规则引擎自动生成当日三餐方案；
3. 展示每餐营养指标与可执行的食材清单、制作步骤。
非核心功能（多日计划、饮食记录、图表看板、云端AI）纳入V2.0。
        """.strip(),
    )

    add_heading(doc, "6. 原型调研与交互验证", 1)
    add_para(doc, "低保真原型采用 900×700 主窗口，三Tab导航，三餐卡片横向排列，详情弹窗展示克数级食材与3步做法。")
    add_para(doc, "关键交互：注册登录 → 目标设定 → 一键推荐 → 查看详情 → 重新生成。")

    add_heading(doc, "7. 调研结论与建议", 1)
    add_multiline(
        doc,
        """
1. 市场需求真实存在，差异化在于「中式菜谱+离线+可解释规则引擎」。
2. 7日工期可行，但必须严格MoSCoW优先级，砍掉记录与图表类功能。
3. 推荐算法采用约束校验优于纯机器学习，便于答辩演示与调试。
4. AI协作应限于CRUD与控件生成，营养阈值与架构决策须人工把关。
        """.strip(),
    )

    out = os.path.join(BASE, f"{STUDENT_ID}_{STUDENT_NAME}_调研报告_V1.0.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    p1 = build_start_report()
    p2 = build_research_report()
    with open(os.path.join(BASE, "generate_log.txt"), "w", encoding="utf-8") as f:
        f.write(f"Start report: {p1}\nResearch report: {p2}\n")
